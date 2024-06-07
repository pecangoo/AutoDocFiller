from datetime import datetime

from docx import Document
import re


class AutoFiller:
    def __init__(self, docx_file):
        self.doc = Document(docx_file)
        self.dict_values = {}

    def date_to_dict(self, tag: str, date_str: str):
        date = datetime.strptime(date_str, '%Y-%m-%d')
        month = str(date.month)
        day = str(date.day)
        if len(month) == 1:
            month = "0" + month
        if len(day) == 1:
            day = "0" + day

        self.dict_values[f"{tag}M"] = month
        self.dict_values[f"{tag}D"] = day
        self.dict_values[f"{tag}Y"] = str(date.year)

    def add_value_to_dict(self, key: str, value: str) -> None:
        self.dict_values[key] = value

    def replace_all_in_docs(self):
        for key in self.dict_values.keys():
            print(key)
            self.replace_text_in_docx(key, self.dict_values[key])

    def replace_text_in_docx(self, old_text, new_text):
        print("\nKey: " + old_text + "\n")

        for paragraph in self.doc.paragraphs:
            if old_text in paragraph.text:
                inline = paragraph.runs
                for i in inline:
                    print(i.text, end=" ::: ")
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        text = inline[i].text.replace(old_text, new_text)
                        inline[i].text = text

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            inline = paragraph.runs
                            # for i in inline:
                            #     print(i.text, end=" ::: ")
                            for i in range(len(inline)):
                                if old_text in inline[i].text:
                                    text = inline[i].text.replace(old_text, new_text)
                                    inline[i].text = text
        for section in self.doc.sections:
            header = section.footer
            for paragraph in header.paragraphs:
                print("Paragraph: " + paragraph.text)
                if old_text in paragraph.text:
                    print("Old_Text: " + old_text)
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if old_text in inline[i].text:
                            text = inline[i].text.replace(old_text, new_text)
                            inline[i].text = text

    def save_doc_to_file(self, path_new_doc):
        self.doc.save(path_new_doc)
